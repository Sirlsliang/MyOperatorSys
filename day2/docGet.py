#!/usr/bin/pyhton3

from docx import Document
from bs4 import BeautifulSoup

from urllib import request

url="http://freebuf.com/news/94263.html"
data = request.urlopen(url)

document = Document()#创建doc 文档

soup = BeautifulSoup(data.read())

article = soup.find(name="div",attrs={'id':'contenttxt'}).children#定位到相关段落
for e in article:#遍历
    try:
        if e.img:
            pic_name=''
            print(e.img.attrs['src'])
            if 'gif' in e.img.attrs['src']:
                pic_name = 'temp.gif'
            elif 'png' in e.img.attrs['src']:
                pic_name = 'tmp.png'
            elif 'jpg' in e.img.attrs['src']:
                pic_name = 'temp.jpg'
            else:
                pic_name = 'temp.jpg'
            request.urlretrieve(e.img.attrs['src'],filename=pic_name)#将远程的数据获取到本地,fileName指定了保存到本地的路径，为指定该选项，urllib会生成一个临时文件文件来保存数据。
            document.add_picture(pic_name)
    except:
        pass
    if e.string:
        print(e.string)
        document.add_paragraph(e.string)
document.save("freebuf_article.doc")
print("Success create a document")
